import docx
import os

def analyzeDocX(inputDocx):

    def analyzeTables():
        tableDocx = doc.tables[0]
        for r in range(0,len(tableDocx.rows)):
            for c in range(0,len(tableDocx.columns)):
                sourceValue = "r"+str(r)+ "c"+str(c)
                targetCell = tableDocx.cell(r, c)
                targetCellValue = targetCell.text
                targetCell.text = str(sourceValue)

        tableDocx = doc.tables[1]
        for r in range(0,len(tableDocx.rows)):
            for c in range(0,len(tableDocx.columns)):
                sourceValue = "r"+str(r)+ "c"+str(c)
                targetCell = tableDocx.cell(r, c)
                print(len(tableDocx.rows))
                targetCellValue = targetCell.text
                #write data  to .docxf file
                targetCell.text = str(sourceValue)
        return

    def analyzeParagraphs():
        """ iterate paragraphs and print listnumber of paragraph to docX"""
        i=0
        for paragraph in doc.paragraphs:
            print(paragraph.text)
            paragraph.text = "paragraph"+str(i)
            i+=1
        return
    def analyzeRun():
        """analyze run and print listnumber of run to docX and to cli"""
        i=0
        for run in paragraph.runs:
            print(run.text, "run: ", i)
            run.text = "run"+str(i)
            i+=1


    #open .docx
    print(inputDocx)
    #outputDocx = "saved-"+inputDocx
    save= "analyze"
    outputDocx = os.path.join(  os.path.expanduser('~'),
                                "schoolReport",
                                "output",
                                 save + docxFile)
    savepath= outputDocx
    doc = docx.Document(inputDocx)

    #analyzeTables()
    #save .docX
    doc.save(savepath)
    print("Was saved in", savepath)




if __name__ == "__main__":

    docxFile = "Z 420 - Notenzeugnis für Schülerinnen und Schüler mit dem Förderbedarf Lernen (01.21)-1.docx"
    inputDocx = os.path.join(os.path.expanduser('~'),"schoolReport", docxFile)

    analyzeDocX(inputDocx)
