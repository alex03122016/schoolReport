import docx
import os

def testwritePupilDetailsToSchoolReport(inputDocx,
                                    docxFile,
                                    nameOfPupil,
                                    Geburtsdatum="",
                                    save= "test"):
    """ will overwrite the inputDocx File
    looks for #Vorname and #Geburtsdatum in docx file and replaces it
    with data given as variable nameOfPupil and Geburtsdatum"""

    #open .docx
    print(inputDocx)
    #outputDocx = "saved-"+inputDocx
    outputDocx = os.path.join(  os.path.expanduser('~'),
                                "schoolReport",
                                "output",
                                 save + docxFile)
    savepath= outputDocx
    doc = docx.Document(inputDocx)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            print(run.text)
            if "#Vorname" in run.text:
                print("OK")
                print(run.text)
                run.text = nameOfPupil

    """for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            print(run.text)
            if "#Geburtsdatum" in run.text:
                print("OK")
                print(run.text)
                run.text = Geburtsdatum"""


    #save .docX
    doc.save(inputDocx)
    print("Was saved in", inputDocx)

    return

if __name__ == "__main__":

    docxFileTest = "Z 420 - Notenzeugnis für Schülerinnen und Schüler mit dem Förderbedarf Lernen (01.21)-1.docx"
    inputDocxTest = os.path.join(os.path.expanduser('~'),"schoolReport", docxFileTest)
    nameOfPupilTest = "Alexander Tanck"
    GeburtsdatumTest = "26. September 1986"
    saveTest = "Test1"

    testwritePupilDetailsToSchoolReport(inputDocx=inputDocxTest,
                                    docxFile=docxFileTest,
                                    nameOfPupil=nameOfPupilTest,
                                    Geburtsdatum=GeburtsdatumTest,
                                    save=saveTest )
