import docx
import os

def writePupilDetailsToSchoolReport(inputDocx,
                                    docxFile,
                                    nameOfPupil,
                                    Geburtsdatum="",
                                    save= "test"):

    #open .docx
    print(inputDocx)
    #outputDocx = "saved-"+inputDocx
    outputDocx = os.path.join(  os.path.expanduser('~'),
                                "schoolReport",
                                "output",
                                 save + docxFile)
    savepath= outputDocx
    doc = docx.Document(inputDocx)


    paragraph = doc.paragraphs[6]
    run= paragraph.runs[2]
    #paragraph 6, run 2
    run.text = nameOfPupil

    """paragraph = doc.paragraphs[8]
    run= paragraph.runs[2]
    #paragraph 6, run 2
    run.text = Geburtsdatum"""

    #save .docX
    doc.save(inputDocx)
    print("Was saved in", inputDocx)

    return

if __name__ == "__main__":

    docxFile = "Z 420 - Notenzeugnis für Schülerinnen und Schüler mit dem Förderbedarf Lernen (01.21)-1.docx"
    inputDocxTest = os.path.join(os.path.expanduser('~'),"schoolReport", docxFile)
    nameOfPupilTest = "Alexander Tanck"
    GeburtsdatumTest = "26. September 1986"
    saveTest = "Test1"

    writePupilDetailsToSchoolReport(inputDocxTest,
                                    nameOfPupilTest,
                                    GeburtsdatumTest,
                                    saveTest )
